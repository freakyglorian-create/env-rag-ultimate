"""
多格式文档解析器 - PDF / Word / HTML / Markdown / TXT / Excel
"""
import re, os
from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings


class DocumentParser:
    SUPPORTED = {".txt", ".md", ".pdf", ".docx", ".doc", ".html", ".htm", ".xlsx"}

    @staticmethod
    def parse_txt(path: str) -> List[Document]:
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        for enc in encodings:
            try:
                with open(path, "r", encoding=enc) as f:
                    content = f.read()
                return [Document(page_content=content, metadata={"source": Path(path).name, "type": "txt"})]
            except (UnicodeDecodeError, Exception):
                continue
        return []

    @staticmethod
    def parse_pdf(path: str) -> List[Document]:
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            docs = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(page_content=text.strip(), metadata={"source": Path(path).name, "page": i, "type": "pdf"}))
            return docs
        except Exception:
            return DocumentParser.parse_txt(path)

    @staticmethod
    def parse_docx(path: str) -> List[Document]:
        try:
            from docx import Document as DX
            doc = DX(path)
            content = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            return [Document(page_content=content, metadata={"source": Path(path).name, "type": "docx"})]
        except Exception:
            return DocumentParser.parse_txt(path)

    @staticmethod
    def parse_html(path: str) -> List[Document]:
        try:
            from bs4 import BeautifulSoup
            with open(path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            for s in soup(["script", "style", "nav", "footer", "header"]):
                s.decompose()
            text = soup.get_text(separator="\n", strip=True)
            return [Document(page_content=text, metadata={"source": Path(path).name, "type": "html"})]
        except Exception:
            return DocumentParser.parse_txt(path)

    @staticmethod
    def parse_markdown(path: str) -> List[Document]:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        content = re.sub(r'#{1,6}\s*', '', content)
        content = re.sub(r'\*\*|__', '', content)
        content = re.sub(r'`{1,3}.*?`{1,3}', '', content, flags=re.DOTALL)
        return [Document(page_content=content, metadata={"source": Path(path).name, "type": "markdown"})]

    @staticmethod
    def parse_xlsx(path: str) -> List[Document]:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True)
            docs = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c else "" for c in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    docs.append(Document(page_content="\n".join(rows), metadata={"source": Path(path).name, "sheet": sheet, "type": "xlsx"}))
            return docs
        except Exception:
            return []

    @classmethod
    def parse(cls, path: str) -> List[Document]:
        ext = Path(path).suffix.lower()
        parsers = {".pdf": cls.parse_pdf, ".docx": cls.parse_docx, ".doc": cls.parse_docx,
                   ".html": cls.parse_html, ".htm": cls.parse_html, ".md": cls.parse_markdown,
                   ".xlsx": cls.parse_xlsx}
        parser = parsers.get(ext, cls.parse_txt)
        return parser(path)

    @classmethod
    def parse_directory(cls, directory: str) -> List[Document]:
        all_docs = []
        for fp in Path(directory).rglob("*"):
            if fp.is_file() and fp.suffix.lower() in cls.SUPPORTED:
                print(f"[Parser] {fp.name}")
                try:
                    all_docs.extend(cls.parse(str(fp)))
                except Exception as e:
                    print(f"[Parser] 失败 {fp.name}: {e}")
        return all_docs

    @staticmethod
    def split_documents(docs: List[Document], chunk_size=None, overlap=None) -> List[Document]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size or settings.CHUNK_SIZE,
            chunk_overlap=overlap or settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "？", "！", ".", "?", "!", "；", ";", " ", ""],
        )
        chunks = splitter.split_documents(docs)
        return [c for c in chunks if len(c.page_content.strip()) > 30]
